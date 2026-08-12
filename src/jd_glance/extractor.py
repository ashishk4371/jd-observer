import re
import io
from typing import List, Dict, Tuple, Optional

# Optional imports for document parsing
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx
except ImportError:
    docx = None


# Expanded list of common technical and professional skills
TECH_SKILLS = [
    # Languages
    "Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "Go", "Golang", "Rust", "Ruby",
    "PHP", "Swift", "Kotlin", "R", "Scala", "Dart", "HTML", "CSS", "SQL", "NoSQL", "Bash", "Shell",

    # Web & Frameworks
    "React", "React Native", "Next.js", "Vue", "Vue.js", "Angular", "Svelte", "Node.js", "Express",
    "FastAPI", "Django", "Flask", "Spring Boot", "ASP.NET", "Laravel", "Tailwind CSS", "Bootstrap",
    "GraphQL", "REST API", "gRPC", "WebSockets", "Microservices",

    # Cloud & DevOps
    "AWS", "Amazon Web Services", "Azure", "GCP", "Google Cloud", "Docker", "Kubernetes", "K8s",
    "Terraform", "Ansible", "CI/CD", "GitHub Actions", "Jenkins", "GitLab CI", "Linux", "Nginx",
    "Helm", "Cloudflare", "Serverless", "Lambda",

    # Data & AI/ML
    "PyTorch", "TensorFlow", "Keras", "Scikit-Learn", "Pandas", "NumPy", "OpenCV", "NLP",
    "LangChain", "LlamaIndex", "LLM", "Generative AI", "Transformers", "BERT", "OpenAI API",
    "PostgreSQL", "MySQL", "MongoDB", "Redis", "Elasticsearch", "Pinecone", "ChromaDB",
    "Snowflake", "BigQuery", "Apache Spark", "Kafka", "Airflow", "ETL", "Databricks",

    # Testing & Tooling
    "Git", "GitHub", "GitLab", "Jira", "Confluence", "Postman", "Jest", "Cypress", "Playwright",
    "Pytest", "JUnit", "Webpack", "Vite", "Docker Compose",

    # Soft & Domain Skills
    "Agile", "Scrum", "Kanban", "System Design", "Object-Oriented Programming", "OOP",
    "Data Structures", "Algorithms", "Product Management", "Leadership", "Code Review",
    "Unit Testing", "TDD", "CI/CD Pipeline", "Problem Solving", "Cross-functional Collaboration"
]

# Build compiled regex patterns for precise skill matching (case-insensitive word boundary)
SKILL_PATTERNS = []
for skill in TECH_SKILLS:
    # Escape special regex characters in skill names like C++, C#, .NET, Node.js
    escaped = re.escape(skill)
    # Ensure word boundaries work for terms with special trailing chars
    pattern = re.compile(rf"(?<!\w){escaped}(?!\w)", re.IGNORECASE)
    SKILL_PATTERNS.append((skill, pattern))


def extract_text_from_pdf(content_bytes: bytes) -> str:
    """Extract clean text from a PDF file using PyMuPDF (fitz)."""
    if not fitz:
        raise ValueError("PyMuPDF is not installed for PDF parsing.")
    doc = fitz.open(stream=content_bytes, filetype="pdf")
    text_chunks = []
    for page in doc:
        text_chunks.append(page.get_text("text"))
    return "\n".join(text_chunks)


def extract_text_from_docx(content_bytes: bytes) -> str:
    """Extract text from a .docx file using python-docx."""
    if not docx:
        raise ValueError("python-docx is not installed for DOCX parsing.")
    doc = docx.Document(io.BytesIO(content_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables if any
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def extract_text_from_file(content_bytes: bytes, filename: str) -> str:
    """Extract text based on file extension (.pdf, .docx, .txt)."""
    fname = filename.lower()
    if fname.endswith(".pdf"):
        return extract_text_from_pdf(content_bytes)
    elif fname.endswith(".docx"):
        return extract_text_from_docx(content_bytes)
    else:
        # Treat as UTF-8 plain text
        return content_bytes.decode("utf-8", errors="ignore")


def extract_skills(text: str) -> List[str]:
    """Extract matching technical and professional skills from text."""
    found_skills = set()
    for skill_name, pattern in SKILL_PATTERNS:
        if pattern.search(text):
            found_skills.add(skill_name)
    return sorted(found_skills)


def extract_years_of_experience(text: str) -> Optional[float]:
    """Extract required or total years of experience using regex heuristics."""
    patterns = [
        r"(\d+(?:\.\d+)?)\s*\+?\s*years?(?:\s+of)?\s+experience",
        r"experience\s*:\s*(\d+(?:\.\d+)?)\s*\+?\s*years?",
        r"minimum\s+(?:of\s+)?(\d+(?:\.\d+)?)\s*\+?\s*years?",
        r"(\d+(?:\.\d+)?)\s*\+?\s*yrs?\s+exp",
    ]
    matches = []
    for pat in patterns:
        for m in re.finditer(pat, text, flags=re.IGNORECASE):
            try:
                matches.append(float(m.group(1)))
            except ValueError:
                pass
    if matches:
        return max(matches)
    return None